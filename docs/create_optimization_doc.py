"""
Script to create a properly formatted Word document for Model Optimization Details.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def set_cell_shading(cell, color):
    """Set background color for a table cell."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def create_table(doc, headers, rows, header_color="1F4E79", widths=None):
    """Create a formatted table with headers and data rows."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # Set header row
    header_row = table.rows[0]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_shading(cell, header_color)
    
    # Set data rows
    for row_idx, row_data in enumerate(rows):
        row = table.rows[row_idx + 1]
        for col_idx, cell_text in enumerate(row_data):
            cell = row.cells[col_idx]
            cell.text = str(cell_text)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            # Alternate row colors
            if row_idx % 2 == 0:
                set_cell_shading(cell, "F2F2F2")
    
    # Set column widths if provided
    if widths:
        for i, width in enumerate(widths):
            for row in table.rows:
                row.cells[i].width = Cm(width)
    
    return table

def add_heading(doc, text, level=1):
    """Add a heading with proper formatting."""
    heading = doc.add_heading(text, level=level)
    if level == 1:
        heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        heading.runs[0].font.size = Pt(16)
    elif level == 2:
        heading.runs[0].font.color.rgb = RGBColor(31, 78, 121)
        heading.runs[0].font.size = Pt(14)
    elif level == 3:
        heading.runs[0].font.color.rgb = RGBColor(68, 114, 196)
        heading.runs[0].font.size = Pt(12)
    return heading

def main():
    # Create document
    doc = Document()
    
    # Set document margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)
    
    # =========================================================================
    # TITLE PAGE
    # =========================================================================
    title = doc.add_heading('Model Optimization Details', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.size = Pt(28)
    title.runs[0].font.color.rgb = RGBColor(31, 78, 121)
    
    subtitle = doc.add_paragraph('LungXAI - Lung Cancer Classification Project')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(14)
    subtitle.runs[0].font.color.rgb = RGBColor(68, 114, 196)
    subtitle.runs[0].italic = True
    
    doc.add_paragraph()
    date_para = doc.add_paragraph('Document Version: 1.0')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_para = doc.add_paragraph('Date: March 2026')
    date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_page_break()
    
    # =========================================================================
    # TABLE OF CONTENTS
    # =========================================================================
    add_heading(doc, 'Table of Contents', 1)
    
    toc_items = [
        ('1.', 'Overview', '3'),
        ('2.', 'Bayesian Optimization (Swin-T)', '4'),
        ('3.', 'Standard Fine-Tuned Models Configuration', '6'),
        ('4.', 'Baseline Models Configuration', '7'),
        ('5.', 'Common Optimization Techniques', '8'),
        ('6.', 'Model Hyperparameters, Ranges & Results', '9'),
        ('  6.1', 'Complete Model Comparison Table', '9'),
        ('  6.2', 'Fine-Tuned CNN Models', '10'),
        ('  6.3', 'Baseline CNN Models', '11'),
        ('  6.4', 'Transformer Models', '13'),
        ('7.', 'Summary Comparison Table', '15'),
        ('8.', 'Recommendations', '16'),
    ]
    
    for num, item, page in toc_items:
        para = doc.add_paragraph()
        para.add_run(f'{num} {item}').bold = True
        para.add_run(f'  {"." * 50}  {page}')
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 1: OVERVIEW
    # =========================================================================
    add_heading(doc, '1. Overview', 1)
    
    doc.add_paragraph(
        'This document provides a comprehensive overview of all optimization techniques '
        'and hyperparameter configurations used across the fine-tuned models in the '
        'LungXAI lung cancer classification project.'
    )
    
    add_heading(doc, '1.1 Models Trained in This Project', 2)
    
    create_table(doc, 
        ['Category', 'Models'],
        [
            ['CNN Models', 'ResNet-50, MobileNetV2, VGG-16, DenseNet-121, EfficientNet-B0'],
            ['Transformer Models', 'ViT-B/16, Swin-T, DeiT-Small, MobileViT-S'],
        ],
        widths=[4, 12]
    )
    
    doc.add_paragraph()
    add_heading(doc, '1.2 Bayesian Optimization Summary', 2)
    
    create_table(doc,
        ['Model', 'Bayesian Optimization', 'Optimization Method'],
        [
            ['Swin-T (Transformer)', '✓ Yes', 'Optuna TPE (30 trials)'],
            ['MobileNetV2', '✗ No', 'Manual hyperparameters'],
            ['ResNet-50', '✗ No', 'Manual hyperparameters'],
            ['ViT-B/16', '✗ No', 'Manual hyperparameters'],
            ['DeiT-Small', '✗ No', 'Manual hyperparameters'],
            ['MobileViT-S', '✗ No', 'Manual hyperparameters'],
            ['DenseNet-121', '✗ No', 'Manual hyperparameters'],
            ['EfficientNet-B0', '✗ No', 'Manual hyperparameters'],
            ['VGG-16', '✗ No', 'Manual hyperparameters'],
        ],
        widths=[5, 4, 6]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 2: BAYESIAN OPTIMIZATION
    # =========================================================================
    add_heading(doc, '2. Bayesian Optimization (Swin-T Only)', 1)
    
    doc.add_paragraph(
        'Only the Swin Transformer (Tiny) model underwent Bayesian hyperparameter '
        'optimization using the Optuna framework with Tree-structured Parzen Estimator (TPE) sampling.'
    )
    
    para = doc.add_paragraph()
    para.add_run('Training Script: ').bold = True
    para.add_run('archive_transformer_files/train_swin_bayesian.py')
    
    add_heading(doc, '2.1 Optimization Configuration', 2)
    
    create_table(doc,
        ['Parameter', 'Value'],
        [
            ['Method', 'Tree-structured Parzen Estimator (TPE)'],
            ['Framework', 'Optuna'],
            ['Number of Trials', '30'],
            ['Epochs per Trial', '15'],
            ['Final Training Epochs', '50'],
            ['Pruning Strategy', 'MedianPruner (n_startup_trials=5, n_warmup_steps=3)'],
            ['Optimization Time', '~328 minutes'],
            ['Final Training Time', '~53 minutes'],
            ['Total Time', '~381 minutes (~6.4 hours)'],
        ],
        widths=[6, 10]
    )
    
    doc.add_paragraph()
    add_heading(doc, '2.2 Hyperparameter Search Space', 2)
    
    create_table(doc,
        ['Hyperparameter', 'Search Range', 'Scale', 'Description'],
        [
            ['Learning Rate', '1e-6 to 1e-3', 'Log', 'Base learning rate for AdamW'],
            ['Weight Decay', '1e-6 to 1e-2', 'Log', 'L2 regularization strength'],
            ['Dropout Rate', '0.1 to 0.7', 'Linear', 'Classifier dropout probability'],
            ['Beta1 (Adam)', '0.85 to 0.99', 'Linear', 'First moment decay rate'],
            ['Beta2 (Adam)', '0.9 to 0.9999', 'Linear', 'Second moment decay rate'],
            ['LR Scheduler Gamma', '0.1 to 0.9', 'Linear', 'LR decay factor'],
            ['LR Scheduler Step', '3 to 10', 'Integer', 'Epochs between LR decay'],
            ['Warmup Epochs', '0 to 5', 'Integer', 'Learning rate warmup period'],
            ['Label Smoothing', '0.0 to 0.2', 'Linear', 'Soft label regularization'],
        ],
        widths=[4, 3, 2, 6]
    )
    
    doc.add_paragraph()
    add_heading(doc, '2.3 Optimal Hyperparameters Found', 2)
    
    create_table(doc,
        ['Hyperparameter', 'Bayesian Optimized', 'Default Value', 'Change'],
        [
            ['Learning Rate', '0.000155', '0.0001', '+55%'],
            ['Weight Decay', '0.000023', '0.0001', '-77%'],
            ['Dropout Rate', '0.333', '0.5', '-33%'],
            ['Beta1 (Adam)', '0.879', '0.9', '-2.3%'],
            ['Beta2 (Adam)', '0.9997', '0.999', '+0.07%'],
            ['LR Scheduler Gamma', '0.70', '0.5', '+40%'],
            ['LR Scheduler Step', '8', '10', '-20%'],
            ['Warmup Epochs', '0', 'N/A', '-'],
            ['Label Smoothing', '0.076', '0.0', 'New'],
        ],
        widths=[4.5, 4, 3.5, 3]
    )
    
    doc.add_paragraph()
    add_heading(doc, '2.4 Key Insights from Bayesian Optimization', 2)
    
    insights = [
        ('Higher Learning Rate:', 'A 55% higher learning rate (0.000155 vs 0.0001) proved beneficial for faster convergence.'),
        ('Lower Regularization:', 'Lower dropout (0.333) and weight decay (0.000023) combined with label smoothing (0.076) as compensation.'),
        ('Gradual LR Decay:', 'Higher gamma (0.70) provides more gradual learning rate decay compared to default (0.5).'),
        ('More Frequent Decay:', 'Step size of 8 (vs 10) provides more frequent LR adjustments during training.'),
    ]
    
    for title, desc in insights:
        para = doc.add_paragraph(style='List Bullet')
        para.add_run(title).bold = True
        para.add_run(f' {desc}')
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 3: STANDARD FINE-TUNED MODELS
    # =========================================================================
    add_heading(doc, '3. Standard Fine-Tuned Models Configuration', 1)
    
    doc.add_paragraph(
        'All models except the Bayesian-optimized Swin-T use the following standard configuration.'
    )
    
    add_heading(doc, '3.1 Training Configuration', 2)
    
    create_table(doc,
        ['Parameter', 'Value', 'Source'],
        [
            ['Optimizer', 'AdamW', 'src/models/model_factory.py'],
            ['Learning Rate', '1e-4 (0.0001)', 'src/utils/config.py'],
            ['Weight Decay', '1e-4 (0.0001)', 'src/utils/config.py'],
            ['Epochs', '50', 'src/utils/config.py'],
            ['Batch Size (CNNs)', '32', 'train_all_models.py'],
            ['Batch Size (ViT)', '8', 'train_all_models.py'],
            ['Batch Size (Swin-T)', '16', 'train_all_models.py'],
            ['Dropout Rate', '0.5', 'src/utils/config.py'],
            ['Early Stopping Patience', '10 epochs', 'src/utils/config.py'],
            ['Pretrained Weights', 'ImageNet', 'All models'],
        ],
        widths=[5, 4, 6]
    )
    
    doc.add_paragraph()
    add_heading(doc, '3.2 Learning Rate Scheduler', 2)
    
    create_table(doc,
        ['Parameter', 'Value'],
        [
            ['Scheduler Type', 'StepLR'],
            ['Step Size', '10 epochs'],
            ['Gamma (Decay Factor)', '0.5'],
        ],
        widths=[6, 6]
    )
    
    doc.add_paragraph()
    add_heading(doc, '3.3 Data Configuration', 2)
    
    create_table(doc,
        ['Parameter', 'Value'],
        [
            ['Image Size', '224 × 224 pixels'],
            ['Train Split', '70%'],
            ['Validation Split', '15%'],
            ['Test Split', '15%'],
            ['Random Seed', '42'],
            ['Loss Function', 'CrossEntropyLoss'],
        ],
        widths=[6, 6]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 4: BASELINE MODELS
    # =========================================================================
    add_heading(doc, '4. Baseline Models Configuration', 1)
    
    doc.add_paragraph(
        'Baseline models are trained from scratch (without pretrained weights) for comparison purposes. '
        'This demonstrates the value of transfer learning.'
    )
    
    add_heading(doc, '4.1 Training Configuration Comparison', 2)
    
    create_table(doc,
        ['Parameter', 'Baseline', 'Fine-Tuned', 'Difference'],
        [
            ['Pretrained Weights', '✗ None', '✓ ImageNet', 'No transfer learning'],
            ['Optimizer', 'Adam', 'AdamW', 'Simpler optimizer'],
            ['Learning Rate', '0.001', '0.0001', '10× higher'],
            ['Weight Decay', '0', '0.0001', 'No L2 regularization'],
            ['Epochs', '30', '50', '20 fewer epochs'],
            ['LR Scheduler', '✗ None', '✓ StepLR', 'No LR decay'],
            ['Early Stopping', '✗ No', '✓ Yes', 'Trains all epochs'],
        ],
        widths=[4.5, 3, 3, 4.5]
    )
    
    doc.add_paragraph()
    add_heading(doc, '4.2 Baseline Models Trained', 2)
    
    create_table(doc,
        ['Model', 'Checkpoint File'],
        [
            ['ResNet-50', 'checkpoints/baseline/best_model_resnet50_baseline.pth'],
            ['MobileNetV2', 'checkpoints/baseline/best_model_mobilenetv2_baseline.pth'],
            ['ViT-B/16', 'checkpoints/baseline/best_model_vit_b_16_baseline.pth'],
            ['Swin-T', 'checkpoints/baseline/best_model_swin_t_baseline.pth'],
            ['DeiT-Small', 'checkpoints/baseline/best_model_deit_small_baseline.pth'],
            ['MobileViT-S', 'checkpoints/baseline/best_model_mobilevit_s_baseline.pth'],
            ['DenseNet-121', 'checkpoints/baseline/best_model_densenet121_baseline.pth'],
            ['EfficientNet-B0', 'checkpoints/baseline/best_model_efficientnet_b0_baseline.pth'],
            ['VGG-16', 'checkpoints/baseline/best_model_vgg16_baseline.pth'],
        ],
        widths=[3.5, 12]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 5: COMMON OPTIMIZATION TECHNIQUES
    # =========================================================================
    add_heading(doc, '5. Common Optimization Techniques', 1)
    
    add_heading(doc, '5.1 Applied to All Models', 2)
    
    create_table(doc,
        ['Technique', 'Description', 'Purpose'],
        [
            ['Transfer Learning', 'ImageNet pretrained weights', 'Leverage learned features'],
            ['Data Augmentation', 'Rotation (±15°), flip, scaling', 'Prevent overfitting'],
            ['Stratified Sampling', 'Maintain class distribution', 'Handle class imbalance'],
            ['Dropout', '0.5 probability in classifier', 'Regularization'],
        ],
        widths=[4, 5.5, 5.5]
    )
    
    doc.add_paragraph()
    add_heading(doc, '5.2 Applied to Fine-Tuned Models Only', 2)
    
    create_table(doc,
        ['Technique', 'Description', 'Purpose'],
        [
            ['Early Stopping', 'Patience of 10 epochs', 'Prevent overfitting'],
            ['StepLR Scheduler', 'Reduce LR by 0.5× every 10 epochs', 'Improve convergence'],
            ['AdamW Optimizer', 'Adam with decoupled weight decay', 'Better generalization'],
            ['Weight Decay', '1e-4 L2 regularization', 'Regularization'],
        ],
        widths=[4, 5.5, 5.5]
    )
    
    doc.add_paragraph()
    add_heading(doc, '5.3 Applied to Bayesian Swin-T Only', 2)
    
    create_table(doc,
        ['Technique', 'Description', 'Purpose'],
        [
            ['Label Smoothing', '0.076 smoothing factor', 'Soft regularization'],
            ['LR Warmup', 'Gradual LR increase', 'Stable training start'],
            ['Adaptive Hyperparameters', 'Optuna TPE optimization', 'Find optimal config'],
            ['Trial Pruning', 'MedianPruner', 'Efficient search'],
        ],
        widths=[4, 5.5, 5.5]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 6: COMPREHENSIVE HYPERPARAMETER TABLE (USER REQUESTED FORMAT)
    # =========================================================================
    add_heading(doc, '6. Model Hyperparameters, Ranges & Results', 1)
    
    doc.add_paragraph(
        'This section provides a comprehensive overview of hyperparameters used for ALL models '
        '(CNN and Transformers), their search ranges (where applicable), and the final results achieved.'
    )
    
    doc.add_paragraph(
        'Table Format: Model | Hyperparameters | Range | Result'
    ).runs[0].italic = True
    
    # =========================================================================
    # MASTER TABLE - ALL MODELS
    # =========================================================================
    add_heading(doc, '6.1 Complete Model Comparison Table', 2)
    
    # Create the main comprehensive table
    create_table(doc,
        ['Model', 'Hyperparameters', 'Range/Value', 'Result'],
        [
            # Fine-Tuned CNN Models
            ['MobileNetV2 (Fine-Tuned)', 'LR, WD, Dropout, Optimizer', 'lr=1e-4, wd=1e-4, drop=0.5, AdamW', 'Acc: 97.40%, F1: 97.40%'],
            ['ResNet-50 (Fine-Tuned)', 'LR, WD, Dropout, Optimizer', 'lr=1e-4, wd=1e-4, drop=0.5, AdamW', 'Acc: 96.97%, F1: 96.95%'],
            
            # Transformer Models (Fine-Tuned)
            ['Swin-T (Bayesian)', 'LR, WD, Dropout, Label Smooth', 'lr=1e-6~1e-3, wd=1e-6~1e-2, drop=0.1~0.7', 'Acc: 97.84%, F1: 97.83%'],
            ['Swin-T (Standard)', 'LR, WD, Dropout, Optimizer', 'lr=1e-4, wd=1e-4, drop=0.5, AdamW', 'Acc: 97.84%, F1: 97.84%'],
            ['ViT-B/16 (Fine-Tuned)', 'LR, WD, Dropout, Optimizer', 'lr=1e-4, wd=1e-4, drop=0.5, AdamW', 'Acc: 93.51%, F1: 93.48%'],
            
            # Baseline Models (No Pretrained)
            ['MobileNetV2 (Baseline)', 'LR, WD, Dropout, Optimizer', 'lr=1e-3, wd=0, drop=0.5, Adam', 'Acc: 89.61%, F1: 89.35%'],
            ['ResNet-50 (Baseline)', 'LR, WD, Dropout, Optimizer', 'lr=1e-3, wd=0, drop=0.5, Adam', 'Acc: 78.79%, F1: 78.98%'],
            ['DenseNet-121 (Baseline)', 'LR, WD, Dropout, Optimizer', 'lr=1e-3, wd=0, drop=0.5, Adam', 'Acc: 84.42%, F1: 82.63%'],
            ['EfficientNet-B0 (Baseline)', 'LR, WD, Dropout, Optimizer', 'lr=1e-3, wd=0, drop=0.5, Adam', 'Acc: 72.29%, F1: 72.62%'],
            ['VGG-16 (Baseline)', 'LR, WD, Dropout, Optimizer', 'lr=1e-3, wd=0, drop=0.5, Adam', 'Acc: 71.43%, F1: 69.39%'],
        ],
        widths=[4.5, 4, 5, 4]
    )
    
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.add_run('Legend: ').bold = True
    para.add_run('LR = Learning Rate, WD = Weight Decay, Acc = Accuracy, F1 = F1-Score')
    
    doc.add_paragraph()
    
    # =========================================================================
    # DETAILED TABLES BY CATEGORY
    # =========================================================================
    add_heading(doc, '6.2 Fine-Tuned CNN Models (Transfer Learning)', 2)
    
    doc.add_paragraph(
        'These models use ImageNet pretrained weights and are fine-tuned on the lung cancer dataset.'
    )
    
    # MobileNetV2 Fine-tuned
    add_heading(doc, 'MobileNetV2 (Fine-Tuned) - Primary Model', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'ImageNet', 'True', '-'],
            ['Optimizer', 'AdamW', 'AdamW', '-'],
            ['Learning Rate', '1e-4 (fixed)', '0.0001', '-'],
            ['Weight Decay', '1e-4 (fixed)', '0.0001', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '50 (early stop)', '50', '-'],
            ['LR Scheduler', 'StepLR', 'step=10, γ=0.5', '-'],
            ['Test Accuracy', '-', '-', '97.40%'],
            ['Precision', '-', '-', '97.50%'],
            ['Recall', '-', '-', '97.40%'],
            ['F1-Score', '-', '-', '97.40%'],
            ['Parameters', '-', '-', '~2.2M'],
            ['Training Time', '-', '-', '~17 min'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    
    # ResNet-50 Fine-tuned
    add_heading(doc, 'ResNet-50 (Fine-Tuned)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'ImageNet', 'True', '-'],
            ['Optimizer', 'AdamW', 'AdamW', '-'],
            ['Learning Rate', '1e-4 (fixed)', '0.0001', '-'],
            ['Weight Decay', '1e-4 (fixed)', '0.0001', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '50 (early stop)', '29 (stopped)', '-'],
            ['LR Scheduler', 'StepLR', 'step=10, γ=0.5', '-'],
            ['Test Accuracy', '-', '-', '96.97%'],
            ['Precision', '-', '-', '96.99%'],
            ['Recall', '-', '-', '96.97%'],
            ['F1-Score', '-', '-', '96.95%'],
            ['Parameters', '-', '-', '~23.5M'],
            ['Training Time', '-', '-', '~7.3 min'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_page_break()
    
    add_heading(doc, '6.3 Baseline CNN Models (Trained from Scratch)', 2)
    
    doc.add_paragraph(
        'These models are trained WITHOUT pretrained weights (random initialization) to demonstrate '
        'the value of transfer learning.'
    )
    
    # MobileNetV2 Baseline
    add_heading(doc, 'MobileNetV2 (Baseline)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'None', 'False', '-'],
            ['Optimizer', 'Adam', 'Adam', '-'],
            ['Learning Rate', '0.001 (fixed)', '0.001', '-'],
            ['Weight Decay', '0 (fixed)', '0', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '30 (fixed)', '30', '-'],
            ['LR Scheduler', 'None', 'None', '-'],
            ['Test Accuracy', '-', '-', '89.61%'],
            ['Precision', '-', '-', '89.86%'],
            ['Recall', '-', '-', '89.61%'],
            ['F1-Score', '-', '-', '89.35%'],
            ['Val Accuracy', '-', '-', '87.83%'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    
    # ResNet-50 Baseline
    add_heading(doc, 'ResNet-50 (Baseline)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'None', 'False', '-'],
            ['Optimizer', 'Adam', 'Adam', '-'],
            ['Learning Rate', '0.001 (fixed)', '0.001', '-'],
            ['Weight Decay', '0 (fixed)', '0', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '30 (fixed)', '30', '-'],
            ['LR Scheduler', 'None', 'None', '-'],
            ['Test Accuracy', '-', '-', '78.79%'],
            ['Precision', '-', '-', '79.36%'],
            ['Recall', '-', '-', '78.79%'],
            ['F1-Score', '-', '-', '78.98%'],
            ['Val Accuracy', '-', '-', '82.17%'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    
    # DenseNet-121 Baseline
    add_heading(doc, 'DenseNet-121 (Baseline)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'None', 'False', '-'],
            ['Optimizer', 'Adam', 'Adam', '-'],
            ['Learning Rate', '0.001 (fixed)', '0.001', '-'],
            ['Weight Decay', '0 (fixed)', '0', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '30 (fixed)', '30', '-'],
            ['LR Scheduler', 'None', 'None', '-'],
            ['Test Accuracy', '-', '-', '84.42%'],
            ['Precision', '-', '-', '85.70%'],
            ['Recall', '-', '-', '84.42%'],
            ['F1-Score', '-', '-', '82.63%'],
            ['Training Time', '-', '-', '~24.4 min'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_page_break()
    
    # EfficientNet-B0 Baseline
    add_heading(doc, 'EfficientNet-B0 (Baseline)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'None', 'False', '-'],
            ['Optimizer', 'Adam', 'Adam', '-'],
            ['Learning Rate', '0.001 (fixed)', '0.001', '-'],
            ['Weight Decay', '0 (fixed)', '0', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '30 (fixed)', '30', '-'],
            ['LR Scheduler', 'None', 'None', '-'],
            ['Test Accuracy', '-', '-', '72.29%'],
            ['Precision', '-', '-', '73.58%'],
            ['Recall', '-', '-', '72.29%'],
            ['F1-Score', '-', '-', '72.62%'],
            ['Training Time', '-', '-', '~21.6 min'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    
    # VGG-16 Baseline
    add_heading(doc, 'VGG-16 (Baseline)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'None', 'False', '-'],
            ['Optimizer', 'Adam', 'Adam', '-'],
            ['Learning Rate', '0.001 (fixed)', '0.001', '-'],
            ['Weight Decay', '0 (fixed)', '0', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '32 (fixed)', '32', '-'],
            ['Epochs', '30 (fixed)', '30', '-'],
            ['LR Scheduler', 'None', 'None', '-'],
            ['Test Accuracy', '-', '-', '71.43%'],
            ['Precision', '-', '-', '69.82%'],
            ['Recall', '-', '-', '71.43%'],
            ['F1-Score', '-', '-', '69.39%'],
            ['Training Time', '-', '-', '~78 min'],
            ['Parameters', '-', '-', '~138M'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    # =========================================================================
    # SECTION 6.4: TRANSFORMER MODELS
    # =========================================================================
    doc.add_page_break()
    
    add_heading(doc, '6.4 Transformer Models (Fine-Tuned)', 2)
    
    doc.add_paragraph(
        'Transformer-based models using ImageNet pretrained weights.'
    )
    
    # Swin-T Bayesian
    add_heading(doc, 'Swin Transformer Tiny (Bayesian Optimized)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Search Range', 'Optimal Value', 'Result'],
        [
            ['Pretrained', 'ImageNet', 'True', '-'],
            ['Optimizer', 'AdamW', 'AdamW', '-'],
            ['Learning Rate', '1e-6 to 1e-3 (log)', '0.000155', '-'],
            ['Weight Decay', '1e-6 to 1e-2 (log)', '0.0000234', '-'],
            ['Dropout Rate', '0.1 to 0.7', '0.333', '-'],
            ['Label Smoothing', '0.0 to 0.2', '0.076', '-'],
            ['Beta1 (Adam)', '0.85 to 0.99', '0.879', '-'],
            ['Beta2 (Adam)', '0.9 to 0.9999', '0.9997', '-'],
            ['LR Scheduler Gamma', '0.1 to 0.9', '0.70', '-'],
            ['LR Scheduler Step', '3 to 10', '8', '-'],
            ['Batch Size', '16 (fixed)', '16', '-'],
            ['Epochs', '50', '50', '-'],
            ['Test Accuracy', '-', '-', '97.84%'],
            ['Precision', '-', '-', '97.86%'],
            ['Recall', '-', '-', '97.84%'],
            ['F1-Score', '-', '-', '97.83%'],
            ['Parameters', '-', '-', '~28M'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    
    # Swin-T Standard
    add_heading(doc, 'Swin Transformer Tiny (Standard)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'ImageNet', 'True', '-'],
            ['Optimizer', 'AdamW', 'AdamW', '-'],
            ['Learning Rate', '1e-4 (fixed)', '0.0001', '-'],
            ['Weight Decay', '1e-4 (fixed)', '0.0001', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '16 (fixed)', '16', '-'],
            ['Epochs', '50 (early stop)', '50', '-'],
            ['LR Scheduler', 'StepLR', 'step=10, γ=0.5', '-'],
            ['Test Accuracy', '-', '-', '97.84%'],
            ['Precision', '-', '-', '97.86%'],
            ['Recall', '-', '-', '97.84%'],
            ['F1-Score', '-', '-', '97.84%'],
            ['Training Time', '-', '-', '~27.7 min'],
            ['Parameters', '-', '-', '~28M'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    
    # ViT-B/16
    add_heading(doc, 'Vision Transformer ViT-B/16 (Fine-Tuned)', 3)
    
    create_table(doc,
        ['Hyperparameter', 'Value/Range', 'Final Value', 'Result'],
        [
            ['Pretrained', 'ImageNet', 'True', '-'],
            ['Optimizer', 'AdamW', 'AdamW', '-'],
            ['Learning Rate', '1e-4 (fixed)', '0.0001', '-'],
            ['Weight Decay', '1e-4 (fixed)', '0.0001', '-'],
            ['Dropout Rate', '0.5 (fixed)', '0.5', '-'],
            ['Batch Size', '8 (fixed)', '8', '-'],
            ['Epochs', '50 (early stop)', '50', '-'],
            ['LR Scheduler', 'StepLR', 'step=10, γ=0.5', '-'],
            ['Test Accuracy', '-', '-', '93.51%'],
            ['Precision', '-', '-', '93.74%'],
            ['Recall', '-', '-', '93.51%'],
            ['F1-Score', '-', '-', '93.48%'],
            ['Training Time', '-', '-', '~79.6 min'],
            ['Parameters', '-', '-', '~86M'],
        ],
        widths=[4, 4, 4, 4]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 7: SUMMARY COMPARISON TABLE
    # =========================================================================
    add_heading(doc, '7. Summary Comparison Table', 1)
    
    add_heading(doc, '7.1 All Models - Results Summary', 2)
    
    create_table(doc,
        ['Model', 'Type', 'Accuracy', 'F1-Score', 'Parameters'],
        [
            ['Swin-T (Bayesian)', 'Transformer', '97.84%', '97.83%', '~28M'],
            ['Swin-T (Standard)', 'Transformer', '97.84%', '97.84%', '~28M'],
            ['MobileNetV2', 'Fine-Tuned CNN', '97.40%', '97.40%', '~2.2M'],
            ['ResNet-50', 'Fine-Tuned CNN', '96.97%', '96.95%', '~23.5M'],
            ['ViT-B/16', 'Transformer', '93.51%', '93.48%', '~86M'],
            ['MobileNetV2', 'Baseline CNN', '89.61%', '89.35%', '~2.2M'],
            ['DenseNet-121', 'Baseline CNN', '84.42%', '82.63%', '~7.0M'],
            ['ResNet-50', 'Baseline CNN', '78.79%', '78.98%', '~23.5M'],
            ['EfficientNet-B0', 'Baseline CNN', '72.29%', '72.62%', '~5.3M'],
            ['VGG-16', 'Baseline CNN', '71.43%', '69.39%', '~138M'],
        ],
        widths=[4, 3, 3, 3, 3]
    )
    
    doc.add_paragraph()
    add_heading(doc, '7.2 Transfer Learning Impact', 2)
    
    create_table(doc,
        ['Model', 'Baseline Accuracy', 'Fine-Tuned Accuracy', 'Improvement'],
        [
            ['ResNet-50', '78.79%', '96.97%', '+18.18%'],
            ['MobileNetV2', '89.61%', '97.40%', '+7.79%'],
        ],
        header_color="2E7D32",
        widths=[4, 4, 4, 4]
    )
    
    doc.add_paragraph()
    add_heading(doc, '7.3 Hyperparameter Configuration Summary', 2)
    
    create_table(doc,
        ['Parameter', 'Fine-Tuned Models', 'Baseline Models'],
        [
            ['Pretrained Weights', 'ImageNet', 'None (Random)'],
            ['Optimizer', 'AdamW', 'Adam'],
            ['Learning Rate', '0.0001', '0.001'],
            ['Weight Decay', '0.0001', '0'],
            ['Dropout Rate', '0.5', '0.5'],
            ['Batch Size', '32', '32'],
            ['Epochs', '50 (early stopping)', '30'],
            ['LR Scheduler', 'StepLR (γ=0.5)', 'None'],
            ['Early Stopping', 'Yes (patience=10)', 'No'],
        ],
        widths=[5, 5.5, 5.5]
    )
    
    doc.add_page_break()
    
    # =========================================================================
    # SECTION 8: RECOMMENDATIONS
    # =========================================================================
    add_heading(doc, '8. Recommendations for Future Work', 1)
    
    add_heading(doc, '8.1 Extend Bayesian Optimization', 2)
    doc.add_paragraph(
        'Apply Optuna-based hyperparameter tuning to other high-performing models '
        '(MobileNetV2, ResNet-50) to potentially improve their performance.'
    )
    
    add_heading(doc, '8.2 Grid Search Alternative', 2)
    doc.add_paragraph(
        'For models where Bayesian optimization is too time-consuming, consider a targeted grid search:'
    )
    
    create_table(doc,
        ['Hyperparameter', 'Suggested Values'],
        [
            ['Learning Rate', '[1e-5, 5e-5, 1e-4, 5e-4]'],
            ['Weight Decay', '[1e-5, 1e-4, 1e-3]'],
            ['Dropout', '[0.3, 0.4, 0.5]'],
        ],
        widths=[5, 10]
    )
    
    doc.add_paragraph()
    add_heading(doc, '8.3 Advanced Techniques to Consider', 2)
    
    techniques = [
        'Mixup/CutMix data augmentation',
        'Cosine annealing with warm restarts',
        'Stochastic Weight Averaging (SWA)',
        'Knowledge distillation from best model',
        'Ensemble methods combining top-3 models',
    ]
    
    for tech in techniques:
        doc.add_paragraph(tech, style='List Bullet')
    
    doc.add_paragraph()
    add_heading(doc, '8.4 Key Files Reference', 2)
    
    create_table(doc,
        ['File', 'Purpose'],
        [
            ['src/utils/config.py', 'Centralized configuration'],
            ['src/models/model_factory.py', 'Model creation and optimizer setup'],
            ['train_all_models.py', 'Multi-model training script'],
            ['train_baseline_models.py', 'Baseline training script'],
            ['archive_transformer_files/train_swin_bayesian.py', 'Bayesian optimization script'],
        ],
        widths=[8, 8]
    )
    
    # =========================================================================
    # FOOTER
    # =========================================================================
    doc.add_paragraph()
    doc.add_paragraph()
    footer = doc.add_paragraph('─' * 60)
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    footer_text = doc.add_paragraph()
    footer_text.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_text.add_run('LungXAI - Lung Cancer Classification with Explainable AI').italic = True
    
    footer_date = doc.add_paragraph('Document Generated: March 2026')
    footer_date.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Save document
    output_path = os.path.join(os.path.dirname(__file__), 'Model_Optimization_Details_Final.docx')
    doc.save(output_path)
    print(f"✓ Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    main()
